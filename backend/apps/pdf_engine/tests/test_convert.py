"""Golden tests for the conversion matrix (phase-06; §10, §18).

The Gotenberg-backed routes (office, HTML, URL) run against the real service
when it is up and are skipped when it is not — they are integration tests of a
separate container, and a suite that cannot run without it is a suite nobody
runs. Everything PyMuPDF does is tested unconditionally.
"""
import io
import zipfile

import fitz
import pytest

from apps.pdf_engine.engine import convert as C
from apps.pdf_engine.exceptions import EngineError, InvalidParams, UnsupportedFileError


def _gotenberg_up() -> bool:
    import requests
    from django.conf import settings

    try:
        response = requests.get(f"{settings.GOTENBERG_URL.rstrip('/')}/health", timeout=3)
        return response.status_code == 200
    except Exception:  # noqa: BLE001
        return False


needs_gotenberg = pytest.mark.skipif(
    not _gotenberg_up(), reason="gotenberg is not running"
)


def _png(color=(255, 0, 0), size=(120, 80)) -> bytes:
    pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, *size))
    pix.set_rect(pix.irect, color)
    return pix.tobytes("png")


def _pages(data: bytes) -> int:
    doc = fitz.open(stream=data, filetype="pdf")
    try:
        return doc.page_count
    finally:
        doc.close()


# --------------------------------------------------------------------------- #
# Routing
# --------------------------------------------------------------------------- #
@pytest.mark.parametrize("filename,kind", [
    ("report.pdf", "pdf"),
    ("photo.JPG", "image"),
    ("scan.tiff", "image"),
    ("page.html", "html"),
    ("letter.docx", "office"),
    ("book.odt", "office"),
    ("sheet.xlsx", "office"),
    ("deck.pptx", "office"),
])
def test_the_import_route_is_chosen_by_extension(filename, kind):
    assert C.kind_of(filename) == kind


@pytest.mark.parametrize("filename", ["virus.exe", "archive.zip", "noextension"])
def test_a_file_we_cannot_convert_says_so(filename):
    with pytest.raises(UnsupportedFileError):
        C.kind_of(filename)


# --------------------------------------------------------------------------- #
# In → PDF (images)
# --------------------------------------------------------------------------- #
def test_a_png_becomes_a_one_page_pdf():
    out, title = C.convert_from(data=_png(), filename="holiday.png")
    assert _pages(out) == 1
    assert title == "holiday"


def test_images_are_fitted_to_a4_by_default():
    out, _ = C.convert_from(data=_png(size=(400, 200)), filename="wide.png")
    doc = fitz.open(stream=out, filetype="pdf")
    try:
        rect = doc[0].rect
    finally:
        doc.close()
    assert round(rect.width) == 595 and round(rect.height) == 842


def test_original_size_makes_the_page_the_shape_of_the_image():
    """"Original" means the page *is* the image rather than a sheet of A4 with
    the image on it. The numbers are points, not pixels: PyMuPDF maps image
    pixels to points at 96 dpi, so a 400×200 px image is a 300×150 pt page —
    its natural printed size, which is what "original" should mean."""
    out = C.images_to_pdf([_png(size=(400, 200))], fit="original")
    doc = fitz.open(stream=out, filetype="pdf")
    try:
        rect = doc[0].rect
    finally:
        doc.close()
    assert rect.width / rect.height == pytest.approx(2.0, abs=0.01)
    assert round(rect.width) != 595, "it was fitted to A4 after all"


def test_a_multipage_tiff_keeps_all_its_frames():
    """phase-06 §Tests: "multipage tiff → page count"."""
    doc = fitz.open()
    for color in ((255, 0, 0), (0, 255, 0), (0, 0, 255)):
        page = doc.new_page(width=200, height=100)
        page.draw_rect(page.rect, color=None, fill=[c / 255 for c in color])
    frames = [doc[i].get_pixmap() for i in range(3)]
    doc.close()

    # PyMuPDF writes single-frame TIFFs, so the multi-frame file is assembled
    # here the way a scanner would produce one.
    tiff = _multiframe_tiff(frames)
    out = C.images_to_pdf([tiff])
    assert _pages(out) == 3


def _multiframe_tiff(pixmaps) -> bytes:
    from PIL import Image

    images = [Image.open(io.BytesIO(p.tobytes("png"))).convert("RGB") for p in pixmaps]
    buf = io.BytesIO()
    images[0].save(buf, format="TIFF", save_all=True, append_images=images[1:])
    return buf.getvalue()


def test_several_images_become_several_pages():
    out = C.images_to_pdf([_png(), _png((0, 0, 255))])
    assert _pages(out) == 2


def test_something_that_is_not_an_image_is_refused():
    with pytest.raises(UnsupportedFileError):
        C.images_to_pdf([b"definitely not an image"])


def test_a_pdf_passed_to_convert_from_is_probed_and_returned(fixture_bytes):
    out, title = C.convert_from(data=fixture_bytes("text.pdf"), filename="report.pdf")
    assert _pages(out) == 3
    assert title == "report"


def test_a_corrupt_pdf_is_caught_by_the_probe(fixture_bytes):
    with pytest.raises(EngineError):
        C.convert_from(data=b"%PDF-1.4 broken", filename="broken.pdf")


# --------------------------------------------------------------------------- #
# PDF → out
# --------------------------------------------------------------------------- #
def test_text_export_has_a_separator_per_page(fixture_bytes):
    blob, name, ctype = C.pdf_to_text(fixture_bytes("text.pdf"))
    text = blob.decode()
    assert name == "document.txt" and ctype.startswith("text/plain")
    assert text.count("--- Page ") == 3
    assert "quick brown fox" in text


def test_image_export_is_a_zip_with_one_file_per_page(fixture_bytes):
    blob, name, ctype = C.pdf_to_images(fixture_bytes("text.pdf"), fmt="png", dpi=72)
    assert name.endswith(".zip") and ctype == "application/zip"
    with zipfile.ZipFile(io.BytesIO(blob)) as archive:
        names = archive.namelist()
        assert names == ["page-0001.png", "page-0002.png", "page-0003.png"]
        assert archive.read(names[0]).startswith(b"\x89PNG")


def test_image_export_honours_the_dpi(fixture_bytes):
    low, _, _ = C.pdf_to_images(fixture_bytes("text.pdf"), dpi=72)
    high, _, _ = C.pdf_to_images(fixture_bytes("text.pdf"), dpi=200)
    assert len(high) > len(low)


@pytest.mark.parametrize("dpi", [10, 1200])
def test_a_silly_dpi_is_refused(fixture_bytes, dpi):
    with pytest.raises(InvalidParams):
        C.pdf_to_images(fixture_bytes("text.pdf"), dpi=dpi)


def test_html_export_is_one_document_not_three(fixture_bytes):
    """PyMuPDF emits a whole HTML document per *page*, so concatenating them
    gives a file with three <html> elements — which no browser parses as the
    author intended."""
    blob, name, ctype = C.pdf_to_html(fixture_bytes("text.pdf"))
    html = blob.decode()
    assert name == "document.html" and ctype.startswith("text/html")
    assert html.count("<html") == 1
    assert html.count("<body") == 1
    assert "quick brown fox" in html


def test_markdown_export_carries_the_text(fixture_bytes):
    blob, name, ctype = C.pdf_to_markdown(fixture_bytes("text.pdf"))
    assert name == "document.md" and ctype.startswith("text/markdown")
    assert "quick brown fox" in blob.decode()


def test_word_export_opens_as_a_docx_and_contains_the_text(fixture_bytes):
    """phase-06 §Tests: "pdf→docx opens via python-docx and contains fixture
    text"."""
    from docx import Document as DocxDocument

    blob, name, ctype = C.pdf_to_docx(fixture_bytes("text.pdf"))
    assert name == "document.docx"
    assert "wordprocessingml" in ctype

    document = DocxDocument(io.BytesIO(blob))
    text = "\n".join(p.text for p in document.paragraphs)
    assert "quick brown fox" in text


def test_pdfa_output_probes_clean_and_claims_conformance(fixture_bytes):
    """phase-06 §Tests: "pdfa output passes pikepdf probe + has XMP `pdfaid`
    marker". Full veraPDF validation is backlog, and saying so is the honest
    version of "PDF/A"."""
    import pikepdf

    blob, name, ctype, report = C.pdf_to_pdfa(fixture_bytes("text.pdf"))
    assert name == "document-pdfa.pdf" and ctype == "application/pdf"
    assert report["claims_pdfa"] is True
    assert C.has_pdfa_marker(blob) is True

    pdf = pikepdf.open(io.BytesIO(blob))
    try:
        assert len(pdf.pages) == 3
    finally:
        pdf.close()


def test_an_unknown_export_format_is_refused(fixture_bytes):
    with pytest.raises(InvalidParams):
        C.convert_to(fixture_bytes("text.pdf"), target="powerpoint")


@pytest.mark.parametrize("target", ["txt", "md", "html", "images"])
def test_every_cheap_export_format_produces_bytes(fixture_bytes, target):
    blob, name, ctype, _ = C.convert_to(fixture_bytes("text.pdf"), target=target, dpi=72)
    assert blob and name and ctype


# --------------------------------------------------------------------------- #
# Gotenberg-backed routes
# --------------------------------------------------------------------------- #
@needs_gotenberg
def test_a_docx_becomes_a_readable_pdf():
    """phase-06 §Tests: "docx→pdf→probe"."""
    from docx import Document as DocxDocument

    document = DocxDocument()
    document.add_paragraph("Converted from Word by ZenPDF.")
    buf = io.BytesIO()
    document.save(buf)

    out, title = C.convert_from(data=buf.getvalue(), filename="letter.docx")
    assert _pages(out) >= 1
    assert title == "letter"
    doc = fitz.open(stream=out, filetype="pdf")
    try:
        assert "Converted from Word" in doc[0].get_text()
    finally:
        doc.close()


@needs_gotenberg
def test_an_html_file_becomes_a_pdf():
    out, _ = C.convert_from(
        data=b"<html><body><h1>Hello from HTML</h1></body></html>",
        filename="page.html",
    )
    doc = fitz.open(stream=out, filetype="pdf")
    try:
        assert "Hello from HTML" in doc[0].get_text()
    finally:
        doc.close()


@pytest.mark.parametrize("url", [
    "http://169.254.169.254/latest/meta-data/",
    "http://localhost:8000/api/documents/",
    "file:///etc/passwd",
    "http://10.0.0.1/",
])
def test_url_conversion_refuses_the_private_space(url):
    """The engine checks as well as the API edge: params are stored on the job
    and could be replayed."""
    with pytest.raises(InvalidParams):
        C.url_to_pdf(url)


@needs_gotenberg
@pytest.mark.parametrize("url", [
    "http://api:8000/api/health/",          # our own API, by service name
    "http://169.254.169.254/latest/meta-data/",
    "file:///etc/passwd",
    # The spellings layer 1 normalises away and layer 2 originally did not.
    # Before this was fixed the container's own log showed Chromium dialling
    # redis and db through the trailing-dot form.
    "http://redis.:6379/",
    "http://db.:5432/",
    "http://169.254.169.254./",
    "http://0177.0.0.1/",
    "http://2130706433/",
])
def test_gotenberg_itself_refuses_the_private_space(url):
    """Layer 2, proven against the running container rather than against a
    string in settings.

    Deliberately bypasses `url_to_pdf`, which would refuse these at layer 1.
    What is being tested is what happens when layer 1 *cannot* help — a
    redirect, or a name that resolved publicly a moment ago — and the answer
    has to be that Chromium will not go there.
    """
    with pytest.raises(EngineError) as exc:
        C._gotenberg("/forms/chromium/convert/url", [], {"url": url},
                     what="Converting that web page")
    assert "conversion_failed" in exc.value.code or "failed" in str(exc.value).lower()
